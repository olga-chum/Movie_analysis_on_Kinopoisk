import seaborn as sns
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from scipy.stats import pearsonr
import pandas as pd


def analyze_ratings_distribution(df, doc):
    # Параметры для анализа оценок Кинопоиска
    rating_kp = df['rating.kp']

    # Среднее, медианное и модальное значения оценок
    mean_rating = rating_kp.mean()
    median_rating = rating_kp.median()
    mode_rating = rating_kp.mode()[0]

    # Проценты высоких и низких оценок
    high_ratings_percentage = (rating_kp > 7).mean() * 100
    low_ratings_percentage = (rating_kp < 5).mean() * 100

    # Добавляем заголовок и текст в документ
    doc.add_heading("Анализ распределения оценок Кинопоиска", level=1)
    doc.add_paragraph(f"Средняя оценка: {mean_rating:.2f}")
    doc.add_paragraph(f"Медианная оценка: {median_rating:.2f}")
    doc.add_paragraph(f"Модальная оценка: {mode_rating:.2f}")
    doc.add_paragraph(f"Процент высоких оценок (больше 7): {high_ratings_percentage:.2f}%")
    doc.add_paragraph(f"Процент низких оценок (меньше 5): {low_ratings_percentage:.2f}%")

    # Настраиваем стиль Seaborn
    sns.set(style="darkgrid", rc={"axes.facecolor": "#1C1C1C", "grid.color": "#333333"})

    # Создание гистограммы
    plt.figure(figsize=(10, 6))
    sns.histplot(df['rating.kp'], bins=20, kde=True, color="#FF6C00", edgecolor="black", alpha=0.7)
    plt.axvline(mean_rating, color="white", linestyle='--', linewidth=2, alpha=0.7,
                label=f'Средняя оценка: {mean_rating:.2f}')
    plt.xlabel('Оценка Кинопоиска', fontsize=14, color='#FF6C00')
    plt.ylabel('Количество фильмов', fontsize=14, color='#FF6C00')
    plt.legend(facecolor="#555555", edgecolor="black", framealpha=0.9,
               fontsize=12, loc="upper left", labelcolor='white')
    plt.xticks(color='#555555')
    plt.yticks(color='#555555')

    # Сохранение графика в изображение
    graph_filename = "ratings_distribution.png"
    plt.savefig(graph_filename, facecolor="#1C1C1C", bbox_inches='tight')
    plt.close()

    # Добавляем изображение графика в документ
    doc.add_heading("Гистограмма распределения оценок", level=2)
    doc.add_picture(graph_filename, width=5000000, height=3000000)

def compare_platform_ratings(df, doc):

    # Рассчитаем коэффициент корреляции Пирсона
    correlation, _ = pearsonr(df['rating.kp'], df['rating.imdb'])

    # Добавляем заголовок и описание
    doc.add_heading("Сравнение оценок Кинопоиска и IMDb", level=1)
    doc.add_paragraph(f"Коэффициент корреляции Пирсона между оценками"
                      f"Кинопоиска и IMDb составляет: {correlation:.2f}")
    doc.add_paragraph("Ниже представлен график сравнения"
                      " оценок на двух платформах, где видна связь между ними:")

    # Задаем стиль для графика
    plt.style.use('dark_background')

    # Создаем график
    plt.figure(figsize=(10, 6))
    sns.scatterplot(
        data=df[df['rating.imdb'] != 0],
        x='rating.kp',
        y='rating.imdb',
        color='orange',
        s=15,
        edgecolor='white',
        alpha=0.1
    )

    # Добавляем трендовую линию
    sns.regplot(
        data=df[df['rating.imdb'] != 0],
        x='rating.kp',
        y='rating.imdb',
        scatter=False,
        color='#FF6C00',
        line_kws={"linewidth": 2, "alpha": 0.8}
    )

    # Добавляем текст с коэффициентом корреляции
    plt.text(6.0, 2.0, f"Коэффициент корреляции: {correlation:.2f}", fontsize=12, color='white')

    # Настройки графика
    plt.title('Сравнение оценок на Кинопоиске и IMDb', fontsize=16, color='white')
    plt.xlabel('Оценка Кинопоиска', fontsize=12, color='white')
    plt.ylabel('Оценка IMDb', fontsize=12, color='white')
    plt.xticks(color='white')
    plt.yticks(color='white')
    plt.grid(True, linestyle='--', alpha=0.3)

    # Устанавливаем одинаковый диапазон для осей X и Y
    plt.xlim(0, 10)
    plt.ylim(0, 10)

    # Сохраняем график в файл изображения
    graph_filename = "comparison_ratings.png"
    plt.savefig(graph_filename, facecolor="#1C1C1C", bbox_inches='tight')
    plt.close()

    # Добавляем график в документ
    doc.add_picture(graph_filename, width=Inches(6))

def analyze_rating_genres(df, doc):

    # Заголовок документа
    doc.add_heading("Анализ средних рейтингов фильмов по жанрам", level=1)
    doc.add_paragraph("В этом разделе представлен анализ средних"
                      " рейтингов фильмов по жанрам на платформах Кинопоиск и IMDb.")

    # Преобразуем строковые списки в реальные списки (если нужно)
    # df['genres'] = df['genres'].apply(ast.literal_eval)

    # Создаем копию датафрейма для анализа по жанрам
    df_genres = df[df['rating.imdb'] != 0].explode('genres').copy()

    # Выбираем только нужные столбцы
    df_genres = df_genres[['genres', 'rating.kp', 'rating.imdb']]

    # Группируем по жанрам и считаем средние значения оценок
    genre_ratings = df_genres.groupby('genres').agg(
        avg_kp_rating=('rating.kp', 'mean'),
        avg_imdb_rating=('rating.imdb', 'mean')
    ).reset_index()

    # Фильтруем жанры, оставляя только те, которые имеют хотя бы 500 фильмов
    genre_counts = df_genres['genres'].value_counts()
    genre_ratings = genre_ratings[genre_ratings['genres'].
    isin(genre_counts[genre_counts >= 500].index)]

    # Настройка стиля
    plt.style.use('dark_background')

    # Преобразуем данные для удобства визуализации
    genre_ratings_long = genre_ratings.melt(id_vars='genres',
                                            value_vars=['avg_kp_rating', 'avg_imdb_rating'])
    genre_ratings_long['variable'] = genre_ratings_long['variable'].replace({
        'avg_kp_rating': 'Кинопоиск',
        'avg_imdb_rating': 'IMDb'
    })

    # Сортируем жанры по значениям среднего рейтинга IMDb
    genre_ratings_long = genre_ratings_long.sort_values(by='value')

    # Построение графика
    plt.figure(figsize=(14, 8))
    sns.barplot(
        data=genre_ratings_long,
        x='value', y='genres', hue='variable',
        palette={'Кинопоиск': '#FF6C00', 'IMDb': '#FFD700'}
    )

    # Настройка легенды
    plt.legend(title='Платформа', loc='upper right', frameon=False)

    # Дополнительные настройки графика
    plt.title('Средний рейтинг фильмов по жанрам на Кинопоиске и IMDb', fontsize=16, color='white')
    plt.xlabel('Средний рейтинг', fontsize=14, color='white')
    plt.ylabel('Жанры', fontsize=14, color='white')
    plt.xticks(color='white', fontsize=12)
    plt.yticks(color='white', fontsize=12)

    # Устанавливаем ограничение по оси X
    plt.xlim(4, None)

    # Добавляем сетку
    plt.grid(True, linestyle='--', alpha=0.3)

    # Сохраняем график в файл изображения
    graph_filename = "genre_ratings_analysis.png"
    plt.savefig(graph_filename, facecolor="#1C1C1C", bbox_inches='tight')
    plt.close()

    # Добавляем изображение в документ
    doc.add_paragraph("График ниже показывает средние оценки фильмов по жанрам:")
    doc.add_picture(graph_filename, width=Inches(6))

def analyze_rating_genres_time(df, doc):

    # Добавляем заголовок
    doc.add_heading("Анализ изменения популярности жанров фильмов", level=1)
    doc.add_paragraph(
        "В данном анализе показано, как менялась популярность"
        " топ-15 жанров фильмов с течением времени. "
        "Данные отображены на графике ниже."
    )

    # Разворачиваем жанры и выбираем нужные столбцы
    df_genres_years = df[df['rating.imdb'] != 0].explode('genres')
    df_genres_years = df_genres_years[['genres', 'year']]

    # Подсчитываем количество фильмов для каждого жанра по годам
    genre_trends = df_genres_years.groupby(['year', 'genres']).size().reset_index(name='count')

    # Получаем топ-15 жанров по общему количеству фильмов
    top_genres = genre_trends['genres'].value_counts().head(15).index

    # Фильтруем данные для топ-15 жанров
    filtered_genre_trends = genre_trends[genre_trends['genres'].isin(top_genres)]

    # Сортируем жанры по их общему количеству фильмов
    genre_order = (filtered_genre_trends.groupby('genres')['count']
                   .sum().sort_values(ascending=False).index)

    # Построение графика
    plt.style.use('dark_background')
    plt.figure(figsize=(15, 8))
    sns.lineplot(
        data=filtered_genre_trends,
        x='year', y='count', hue='genres', palette='tab10', lw=2, hue_order=genre_order
    )

    # Настройка графика
    plt.title('Изменение популярности топ-15 жанров фильмов'
              ' с течением времени', fontsize=16, color='white')
    plt.xlabel('Год', fontsize=12, color='white')
    plt.ylabel('Количество фильмов', fontsize=12, color='white')
    plt.xticks(color='white', fontsize=10)
    plt.yticks(color='white', fontsize=10)
    plt.grid(True, linestyle='--', alpha=0.3)
    plt.legend(title='Жанры', loc='upper left', bbox_to_anchor=(1, 1), frameon=False)

    # Сохраняем график в изображение
    graph_filename = "genre_trends_over_time.png"
    plt.tight_layout()
    plt.savefig(graph_filename, facecolor="#1C1C1C", bbox_inches='tight')
    plt.close()

    # Добавляем описание и изображение графика в документ Word
    doc.add_paragraph("На графике ниже представлено изменение популярности топ-15 жанров:")
    doc.add_picture(graph_filename, width=Inches(6))

def analyze_rating_genres_trends(df, doc):
    doc.add_heading("Изменение популярности топ-15 жанров фильмов с 2000 по 2020 год", level=1)
    doc.add_paragraph(
        "В этом анализе представлена динамика изменения"
        " количества фильмов для топ-15 жанров "
        "за период с 2000 по 2020 год. График ниже демонстрирует"
        " распределение популярности жанров по годам."
    )

    # Разворачиваем жанры и выбираем нужные столбцы
    df_genres_years = df[df['rating.imdb'] != 0].explode('genres')
    df_genres_years = df_genres_years[['genres', 'year']]

    # Подсчитываем количество фильмов для каждого жанра по годам
    genre_trends = df_genres_years.groupby(['year', 'genres']).size().reset_index(name='count')

    # Получаем топ-15 жанров по общему количеству фильмов
    top_genres = genre_trends['genres'].value_counts().head(15).index

    # Фильтруем данные для топ-15 жанров и нужного временного диапазона
    filtered_genre_trends = genre_trends[
        (genre_trends['genres'].isin(top_genres)) &
        (genre_trends['year'] >= 2000) &
        (genre_trends['year'] <= 2020)
    ]

    # Сортируем жанры по общему количеству фильмов
    genre_order = (filtered_genre_trends.groupby('genres')['count']
                   .sum().sort_values(ascending=False).index)

    # Построение графика
    plt.style.use('dark_background')
    plt.figure(figsize=(9, 6))
    sns.lineplot(
        data=filtered_genre_trends,
        x='year', y='count', hue='genres', palette='tab10', lw=2, hue_order=genre_order
    )

    # Настройка графика
    plt.title('Изменение популярности топ-15 жанров фильмов'
              ' с 2000 по 2020 год', fontsize=16, color='white')
    plt.xlabel('Год', fontsize=12, color='white')
    plt.ylabel('Количество фильмов', fontsize=12, color='white')
    plt.xticks(color='white', fontsize=10)
    plt.yticks(color='white', fontsize=10)
    plt.grid(True, linestyle='--', alpha=0.3)
    plt.legend(title='Жанры', loc='upper left', bbox_to_anchor=(1, 1), frameon=False)

    # Сохраняем график как изображение
    graph_filename = "top_genre_trends.png"
    plt.tight_layout()
    plt.savefig(graph_filename, facecolor="#1C1C1C", bbox_inches='tight')
    plt.close()

    # Добавляем описание и изображение в документ Word
    doc.add_paragraph("График изменения популярности жанров:")
    doc.add_picture(graph_filename, width=Inches(6))

def analyze_budgets(df, doc):
    doc.add_heading("Анализ бюджетов и мировых сборов фильмов", level=1)
    doc.add_paragraph(
        "В данном анализе представлены данные о фильмах с"
        " их бюджетами и мировыми сборами. "
        "Фильтруются фильмы с минимальным бюджетом и сборами, а"
        " также с количеством голосов больше 1000."
    )

    # Фильтруем строки, где бюджет и сборы слишком маленькие или нулевые
    df_budget_fees = df[
        (df['budget_rub'] > 50000) &
        (df['fees_rub_world'] > 50000) &
        (df['votes.kp'] > 1000)
    ]

    # Добавляем отфильтрованные данные в документ Word
    doc.add_heading("Отфильтрованные данные (первые строки):", level=2)
    filtered_data_table = df_budget_fees[['budget_rub', 'fees_rub_world', 'votes.kp']].head()

    # Добавляем таблицу с первыми строками
    table = doc.add_table(rows=1, cols=len(filtered_data_table.columns))
    table.style = 'Table Grid'

    # Добавляем заголовки столбцов
    hdr_cells = table.rows[0].cells
    for idx, column_name in enumerate(filtered_data_table.columns):
        hdr_cells[idx].text = column_name

    # Добавляем строки с данными
    for _, row in filtered_data_table.iterrows():
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = str(value)

    # Строим график пузырьков
    plt.figure(figsize=(10, 6))
    sns.scatterplot(
        x='budget_rub',
        y='fees_rub_world',
        size='votes.kp',  # Размер пузырьков зависит от числа голосов
        sizes=(20, 200),  # Диапазон размеров пузырьков
        data=df_budget_fees,
        alpha=0.6,  # Полупрозрачность пузырьков
        legend=None
    )

    # Настройка логарифмической шкалы
    plt.xscale('log')
    plt.yscale('log')

    # Настройка заголовков и подписей
    plt.title('График пузырьков: Бюджет фильмов и мировые сборы (фильтр выбросов)', fontsize=16)
    plt.xlabel('Бюджет', fontsize=12)
    plt.ylabel('Мировые сборы', fontsize=12)

    # Сохраняем график как изображение
    graph_filename = "bubble_chart_budget_vs_fees.png"
    plt.tight_layout()
    plt.savefig(graph_filename, bbox_inches='tight', facecolor="white")
    plt.close()

    # Добавляем график в документ Word
    doc.add_heading("График пузырьков:", level=2)
    doc.add_paragraph("На графике представлен анализ зависимости"
                      " между бюджетами фильмов и их мировыми сборами.")
    doc.add_picture(graph_filename, width=Inches(6))


def analyze_budgets_and_fees(df, doc):
    pd.set_option('display.max_columns', None)  # Показывать все столбцы
    pd.set_option('display.width', 1000)  # Увеличить ширину вывода

    doc.add_heading("Анализ бюджетов и сборов фильмов", level=1)

    # Фильтруем строки, где бюджет и сборы слишком маленькие или нулевые
    df_budget_fees = df[
        (df['budget_rub'] > 50000) &
        (df['fees_rub_world'] > 50000) &
        (df['votes.kp'] > 1000)
    ].copy()

    # Добавляем новый столбец для разницы между сборами и бюджетом
    df_budget_fees.loc[:, 'fee_budget_diff'] = df_budget_fees['fees_rub_world'] - df_budget_fees['budget_rub']

    # Категории фильмов
    high_budget_high_fees = df_budget_fees[
        (df_budget_fees['budget_rub'] > 1e8) & (df_budget_fees['fees_rub_world'] > 1e9)
    ].sort_values(by='fee_budget_diff', ascending=True)

    high_budget_low_fees = df_budget_fees[
        (df_budget_fees['budget_rub'] > 1e8) & (df_budget_fees['fees_rub_world'] < 1e6)
    ].sort_values(by='fee_budget_diff', ascending=False)

    low_budget_high_fees = df_budget_fees[
        (df_budget_fees['budget_rub'] < 1e6) & (df_budget_fees['fees_rub_world'] > 1e7)
    ].sort_values(by='fee_budget_diff', ascending=False)

    low_budget_low_fees = df_budget_fees[
        (df_budget_fees['budget_rub'] < 1e6) & (df_budget_fees['fees_rub_world'] < 1e6)
    ].sort_values(by='fee_budget_diff', ascending=True)

    # Функция для добавления данных в Word-документ
    def add_movies_to_doc(title, data, doc):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(data.columns))
        table.style = 'Table Grid'

        # Добавляем заголовки
        hdr_cells = table.rows[0].cells
        for idx, column_name in enumerate(data.columns):
            hdr_cells[idx].text = column_name

        # Добавляем строки
        for _, row in data.iterrows():
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = str(value)

        doc.add_paragraph()  # Пустая строка

    # Добавляем категории фильмов в документ
    add_movies_to_doc("Фильмы с большим бюджетом и большими сборами:",
                      high_budget_high_fees[['name', 'genres', 'year', 'budget_rub',
                                             'fees_rub_world', 'rating.kp', 'rating.imdb']].head(5), doc)

    add_movies_to_doc("Фильмы с большим бюджетом, но маленькими сборами:",
                      high_budget_low_fees[['name', 'genres', 'year', 'budget_rub',
                                            'fees_rub_world', 'rating.kp', 'rating.imdb']].head(5), doc)

    add_movies_to_doc("Фильмы с маленьким бюджетом, но большими сборами:",
                      low_budget_high_fees[['name', 'genres', 'year', 'budget_rub',
                                            'fees_rub_world', 'rating.kp', 'rating.imdb']].head(5), doc)

    add_movies_to_doc("Фильмы с маленьким бюджетом и маленькими сборами:",
       low_budget_low_fees[['name', 'genres', 'year', 'budget_rub',
       'fees_rub_world', 'rating.kp', 'rating.imdb']].head(5), doc)


def analyze_top_persons(df, doc):
    doc.add_heading("Анализ лучших актёров и режиссёров", level=1)

    # Фильтруем фильмы с высокими оценками
    high_rating_movies = df[(df['rating.kp'] > 7.5) | (df['rating.imdb'] > 7.5)]

    # Разделяем список актёров и режиссёров
    actors_df = high_rating_movies.explode('actors')
    directors_df = high_rating_movies.explode('directors')

    # Подсчитываем средние рейтинги фильмов для каждого актёра
    actor_ratings = actors_df.groupby('actors').agg(
        avg_kp_rating=('rating.kp', 'mean'),
        avg_imdb_rating=('rating.imdb', 'mean')
    ).reset_index()

    # Сортируем актёров по среднему рейтингу
    actor_ratings_sorted = actor_ratings.sort_values(by=['avg_kp_rating',
        'avg_imdb_rating'], ascending=False)

    # Топ-10 актёров с самыми высокими рейтингами
    top_actors_high_rating = actor_ratings_sorted.head(10)
    top_actors_high_rating.columns = ['actor', 'avg_kp_rating', 'avg_imdb_rating']

    # Подсчитываем средние рейтинги фильмов для каждого режиссёра
    director_ratings = directors_df.groupby('directors').agg(
        avg_kp_rating=('rating.kp', 'mean'),
        avg_imdb_rating=('rating.imdb', 'mean')
    ).reset_index()

    # Сортируем режиссёров по среднему рейтингу
    director_ratings_sorted = director_ratings.sort_values(by=['avg_kp_rating', 'avg_imdb_rating'],
                                                           ascending=False)

    # Топ-10 режиссёров с самыми высокими рейтингами
    top_directors_high_rating = director_ratings_sorted.head(10)
    top_directors_high_rating.columns = ['director', 'avg_kp_rating', 'avg_imdb_rating']

    # Функция для добавления таблицы в Word
    def add_table_to_doc(title, dataframe, doc):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Table Grid'

        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        for idx, column_name in enumerate(dataframe.columns):
            hdr_cells[idx].text = column_name

        # Заполнение таблицы данными
        for _, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = str(round(value, 2) if isinstance(value, (float, int)) else value)

        doc.add_paragraph()  # Пустая строка для разделения

    # Добавляем топ-10 актёров
    add_table_to_doc("Топ-10 актёров с самыми высокими рейтингами:",
                     top_actors_high_rating, doc)

    # Добавляем топ-10 режиссёров
    add_table_to_doc("Топ-10 режиссёров с самыми высокими рейтингами:",
                     top_directors_high_rating, doc)


def analyze_low_persons(df, doc):
    doc.add_heading("Анализ актёров и режиссёров с низкими рейтингами", level=1)

    # Фильтруем фильмы с низкими оценками
    low_rating_movies = df[(df['rating.kp'] < 5.5) | (df['rating.imdb'] < 5.5)]

    # Разделяем список актёров и режиссёров
    actors_df = low_rating_movies.explode('actors')
    directors_df = low_rating_movies.explode('directors')

    # Подсчитываем средние рейтинги фильмов для каждого актёра
    actor_ratings = actors_df.groupby('actors').agg(
        avg_kp_rating=('rating.kp', 'mean'),
        avg_imdb_rating=('rating.imdb', 'mean')
    ).reset_index()

    # Сортируем актёров по среднему рейтингу от низкого к высокому
    actor_ratings_sorted = actor_ratings.sort_values(by=['avg_kp_rating', 'avg_imdb_rating'], ascending=True)

    # Топ-10 актёров с самыми низкими рейтингами
    top_actors_low_rating = actor_ratings_sorted.head(10)
    top_actors_low_rating.columns = ['actor', 'avg_kp_rating', 'avg_imdb_rating']

    # Подсчитываем средние рейтинги фильмов для каждого режиссёра
    director_ratings = directors_df.groupby('directors').agg(
        avg_kp_rating=('rating.kp', 'mean'),
        avg_imdb_rating=('rating.imdb', 'mean')
    ).reset_index()

    # Сортируем режиссёров по среднему рейтингу от низкого к высокому
    director_ratings_sorted = director_ratings.sort_values(by=['avg_kp_rating',
                                                               'avg_imdb_rating'], ascending=True)

    # Топ-10 режиссёров с самыми низкими рейтингами
    top_directors_low_rating = director_ratings_sorted.head(10)
    top_directors_low_rating.columns = ['director', 'avg_kp_rating', 'avg_imdb_rating']

    # Функция для добавления таблицы в Word
    def add_table_to_doc(title, dataframe, doc):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Table Grid'

        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        for idx, column_name in enumerate(dataframe.columns):
            hdr_cells[idx].text = column_name

        # Заполнение таблицы данными
        for _, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = str(round(value, 2) if isinstance(value, (float, int)) else value)

        doc.add_paragraph()  # Пустая строка для разделения

    # Добавляем топ-10 актёров с низкими рейтингами
    add_table_to_doc("Топ-10 актёров с самыми низкими рейтингами:",
                     top_actors_low_rating, doc)

    # Добавляем топ-10 режиссёров с низкими рейтингами
    add_table_to_doc("Топ-10 режиссёров с самыми низкими рейтингами:",
                     top_directors_low_rating, doc)

def analyze_all(df):
    doc = Document()

    # test_analyze_ratings_distribution
    analyze_ratings_distribution(df, doc)

    # test_compare_platform_ratings
    compare_platform_ratings(df, doc)

    # test_analyze_rating_genres
    analyze_rating_genres(df, doc)

    # test_analyze_rating_genres_time
    analyze_rating_genres_time(df, doc)

    # test_analyze_rating_genres_trends
    analyze_rating_genres_trends(df, doc)

    # test_analyze_budgets
    analyze_budgets(df, doc)

    # test_analyze_budgets_and_fees
    analyze_budgets_and_fees(df, doc)

    # test_analyze_top_persons
    analyze_top_persons(df, doc)

    # test_analyze_low_persons
    analyze_low_persons(df, doc)

    doc.save("analysis_result.docx")